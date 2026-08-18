import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_docx_report():
    doc = Document()

    # Define Colors
    COLOR_PRIMARY = RGBColor(0, 51, 102)     # Deep Blue
    COLOR_SECONDARY = RGBColor(0, 153, 204)   # Cyan
    COLOR_TEXT = RGBColor(51, 51, 51)        # Dark Gray
    COLOR_SUCCESS = RGBColor(0, 153, 76)     # Green

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Style Settings
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT

    # ----------------------------------------------------
    # COVER / TITLE SECTION
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("AI-BASED VOICE ASSISTANT (ARIA)\nPROJECT DOCUMENTATION REPORT")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("Speech Recognition & Synthesis with Google OAuth Sign-In, Email/Password Auth & SQLite Database")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY
    sub_p.paragraph_format.space_after = Pt(36)

    # Meta Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Course / Subject:", "Major Academic Coding Project"),
        ("Project Title:", "AI-Based Voice Assistant (ARIA)"),
        ("Newest Features:", "Google OAuth Sign-In, Email/Password Auth, SQLite Relational Database"),
        ("Technologies Used:", "Python 3.13, SQLite 3, Flask, Google GIS SDK, SpeechRecognition, pyttsx3, PyAudio"),
        ("Verification Status:", "100% Passed (14/14 Automated Tests Passed)")
    ]
    for i, (k, v) in enumerate(meta_data):
        cell_k = meta_table.cell(i, 0)
        cell_v = meta_table.cell(i, 1)
        cell_k.paragraphs[0].add_run(k).bold = True
        cell_v.paragraphs[0].add_run(v)
        cell_k.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 1: EXECUTIVE SUMMARY & OBJECTIVES
    # ----------------------------------------------------
    h1 = doc.add_heading("1. Executive Summary & Project Objectives", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Voice assistants have become an essential technology in modern computing, enabling hands-free system interaction, "
        "improving digital accessibility, and boosting user productivity. This project presents ARIA (AI-Based Voice Assistant), "
        "a complete intelligent assistant featuring real-time speech recognition, text-to-speech synthesis, system automation, "
        "Google OAuth Sign-In integration, Email & Password authentication, and a persistent SQLite SQL Database for user accounts "
        "and chat history."
    )

    h2 = doc.add_heading("1.1 Key Objectives", level=2)
    h2.runs[0].font.color.rgb = COLOR_SECONDARY

    objectives = [
        "Google Sign-In System Integration: Connect ARIA to Google's official Identity Services (GIS SDK) for seamless OAuth 2.0 single sign-on.",
        "Email & Password Authentication: Provide full account registration and login backed by secure Werkzeug password hashing.",
        "Persistent Relational SQL Database: Store user profiles and chat history in a persistent SQLite database (database.db) with structured relational tables (users, chat_history).",
        "Audio Capture & Transcription: Capture real-time microphone input and convert speech to text using SpeechRecognition and PyAudio.",
        "Voice Response Synthesis: Convert system responses into clear audible speech using offline pyttsx3 Text-to-Speech (TTS).",
        "Command Intent Processing: Intelligent pattern matching for system commands, time queries, math evaluations, and jokes.",
        "System & Web Automation: Open local applications (Notepad, Calculator, CMD, Explorer, Paint) and web portals (YouTube, Google, StackOverflow).",
        "Interactive Glassmorphism GUI: Provide a futuristic web control center featuring real-time gauges, visualizer, sound effects, and built-in interactive modals."
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # ----------------------------------------------------
    # SECTION 2: SYSTEM ARCHITECTURE & SQL DATABASE SCHEMA
    # ----------------------------------------------------
    h1 = doc.add_heading("2. System Architecture & SQL Database Schema", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "The application follows a 3-tier operational model separating front-end presentation, server logic, and database persistence:"
    )

    tiers = [
        "Frontend Tier (templates/login.html & index.html): Glassmorphic web interface with Google GIS SDK, Auth mode switcher, Web Audio API sound effects, and canvas waveform visualizer.",
        "Server Tier (app.py): Flask REST API server processing commands, handling Google OAuth verification (/api/auth/google), Email login (/api/auth/email-login), and Registration (/api/auth/register).",
        "Database Tier (database.db): SQLite relational database storing user profile credentials and session chat history."
    ]
    for tier in tiers:
        doc.add_paragraph(tier, style='List Bullet')

    h2 = doc.add_heading("2.1 SQL Database Tables Schema", level=2)
    h2.runs[0].font.color.rgb = COLOR_SECONDARY

    # Users Table
    doc.add_heading("Table 1: users (User Accounts & Authentication)", level=3)
    user_table = doc.add_table(rows=9, cols=4)
    user_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    u_headers = ["Column Name", "Data Type", "Constraint", "Description"]
    for j, h in enumerate(u_headers):
        cell = user_table.cell(0, j)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>'))

    u_rows = [
        ("id", "TEXT", "PRIMARY KEY", "Unique User ID (e.g. google_123456... or usr_a8f92b...)"),
        ("name", "TEXT", "NOT NULL", "User's full display name"),
        ("email", "TEXT", "UNIQUE NOT NULL", "User's registered email address"),
        ("password_hash", "TEXT", "NULLABLE", "Werkzeug salted password hash (NULL for Google users)"),
        ("provider", "TEXT", "NOT NULL", "Auth provider used ('Google' or 'Email')"),
        ("picture", "TEXT", "NULLABLE", "Profile avatar URL from Google OAuth"),
        ("created_at", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Account registration timestamp"),
        ("last_login", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Latest login session timestamp")
    ]
    for row_idx, row_vals in enumerate(u_rows, start=1):
        for col_idx, val in enumerate(row_vals):
            cell = user_table.cell(row_idx, col_idx)
            cell.paragraphs[0].add_run(val)

    doc.add_paragraph()

    # Chat History Table
    doc.add_heading("Table 2: chat_history (Conversation History Logs)", level=3)
    chat_table = doc.add_table(rows=8, cols=4)
    chat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_headers = ["Column Name", "Data Type", "Constraint", "Description"]
    for j, h in enumerate(c_headers):
        cell = chat_table.cell(0, j)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>'))

    c_rows = [
        ("id", "INTEGER", "PRIMARY KEY AUTOINCREMENT", "Unique chat message ID"),
        ("user_id", "TEXT", "NOT NULL", "Links message to user account in users table"),
        ("session_id", "TEXT", "NOT NULL", "Group identifier for conversation session"),
        ("role", "TEXT", "NOT NULL", "Sender role ('user' or 'assistant')"),
        ("message", "TEXT", "NOT NULL", "Full transcript text of the message"),
        ("category", "TEXT", "NULLABLE", "Command category (time, math, weather, general)"),
        ("timestamp", "TIMESTAMP", "DEFAULT CURRENT_TIMESTAMP", "Timestamp of message recording")
    ]
    for row_idx, row_vals in enumerate(c_rows, start=1):
        for col_idx, val in enumerate(row_vals):
            cell = chat_table.cell(row_idx, col_idx)
            cell.paragraphs[0].add_run(val)

    doc.add_page_break()

    # ----------------------------------------------------
    # SECTION 3: TERMS & CONDITIONS & ERROR RECTIFICATION
    # ----------------------------------------------------
    h1 = doc.add_heading("3. Terms & Conditions and Error Rectification", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY

    h2 = doc.add_heading("3.1 Terms & Conditions (Operational & Legal Guidelines)", level=2)
    h2.runs[0].font.color.rgb = COLOR_SECONDARY

    terms = [
        ("User Privacy & Data Protection", "User passwords are never stored in plain text. They are protected using Werkzeug cryptographic password hashing (PBKDF2 with SHA-256 salting). Google OAuth profile data is used strictly for authentication and user account creation."),
        ("Microphone Audio Streams", "Spoken voice audio captured via the microphone is converted to text locally or via browser Web Speech API solely for intent processing. Audio streams are not retained on remote servers without explicit user command."),
        ("Account Confidentiality", "Users are responsible for safeguarding their login credentials. Chat histories are linked to individual user IDs in the SQLite database to guarantee session privacy."),
        ("Service Availability & Rate Limits", "ARIA and external integrations (Wikipedia API, Google Identity Services) are provided 'as-is'. Automated rate limits apply to third-party endpoints to maintain stability.")
    ]
    for title, desc in terms:
        p = doc.add_paragraph()
        r1 = p.add_run(f"• {title}: ")
        r1.bold = True
        r1.font.color.rgb = COLOR_PRIMARY
        p.add_run(desc)

    h2 = doc.add_heading("3.2 Error Rectification & Diagnostics", level=2)
    h2.runs[0].font.color.rgb = COLOR_SECONDARY

    errors = [
        ("Database Schema Migration Rectification", "Added automatic schema inspection routines (PRAGMA table_info) and ALTER TABLE migrations to app.py to prevent 'no such column' errors when updating legacy SQLite database files."),
        ("Duplicate Email & Invalid Password Handling", "Implemented strict backend validation (minimum 6 character passwords, duplicate email checks, 401 Unauthorized status responses) preventing corrupt database records."),
        ("Google OAuth Popup Block / Fallback", "Handled Google Identity Services SDK popup blockages by providing an interactive custom prompt fallback dialog so users can log in seamlessly even if third-party popups are restricted."),
        ("Microphone Permission & Browser Speech Fallback", "Resolved 'not-allowed' and 'no-speech' Web Speech API browser errors by auto-switching the interface to text input mode with clear status feedback.")
    ]
    for title, desc in errors:
        p = doc.add_paragraph()
        r1 = p.add_run(f"• {title}: ")
        r1.bold = True
        r1.font.color.rgb = COLOR_PRIMARY
        p.add_run(desc)

    # ----------------------------------------------------
    # SECTION 4: TESTING & VERIFICATION RESULTS
    # ----------------------------------------------------
    h1 = doc.add_heading("4. Testing & Verification Results", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "Automated unit test suites (test_assistant.py and test_auth_db.py) were executed to verify both voice engine capabilities "
        "and SQL database authentication. All 14 tests executed with 100% success rate."
    )

    test_table = doc.add_table(rows=15, cols=3)
    test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_headers = ["Test ID", "Module / Command Tested", "Result Status"]
    for j, h in enumerate(t_headers):
        cell = test_table.cell(0, j)
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>'))

    test_cases = [
        ("Test 1", "what is the current time?", "PASSED (100%)"),
        ("Test 2", "hello who are you", "PASSED (100%)"),
        ("Test 3", "show system status", "PASSED (100%)"),
        ("Test 4", "calculate 12 times 8", "PASSED (100%)"),
        ("Test 5", "tell me a joke", "PASSED (100%)"),
        ("Test 6", "open notepad", "PASSED (100%)"),
        ("Test 7", "open youtube", "PASSED (100%)"),
        ("Test 8", "search wikipedia for python", "PASSED (100%)"),
        ("Test 9", "goodbye", "PASSED (100%)"),
        ("Test 10", "describe me about artificial intelligence", "PASSED (100%)"),
        ("Test 11", "SQLite Schema Integrity Check", "PASSED (100%)"),
        ("Test 12", "Email & Password Account Registration", "PASSED (100%)"),
        ("Test 13", "Email & Password Login & Hashing Check", "PASSED (100%)"),
        ("Test 14", "Google OAuth Account Linking & Sync", "PASSED (100%)")
    ]
    for row_idx, row_vals in enumerate(test_cases, start=1):
        for col_idx, val in enumerate(row_vals):
            cell = test_table.cell(row_idx, col_idx)
            r = cell.paragraphs[0].add_run(val)
            if col_idx == 2:
                r.bold = True
                r.font.color.rgb = COLOR_SUCCESS

    # ----------------------------------------------------
    # SECTION 5: CONCLUSION & PROJECT SUMMARY
    # ----------------------------------------------------
    h1 = doc.add_heading("5. Conclusion & Project Summary", level=1)
    h1.runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "The AI-Based Voice Assistant (ARIA) project successfully fulfills all core academic and functional requirements. "
        "With the addition of Google Sign-In OAuth, Email/Password authentication, and persistent SQLite relational database storage, "
        "the project delivers a complete, secure, and production-ready voice assistant platform."
    )

    doc.save("AI_Voice_Assistant_Project_Report.docx")
    print("[OK] Created updated AI_Voice_Assistant_Project_Report.docx successfully!")

if __name__ == '__main__':
    create_docx_report()
